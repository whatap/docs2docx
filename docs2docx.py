import requests
from bs4 import BeautifulSoup, NavigableString
from bs4.element import Tag
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image as PILImage
import re
import os
import base64
import sys
import argparse
from docx.enum.style import WD_STYLE_TYPE

def clean_text(text):
    """주어진 텍스트에서 특수 문자를 제거하고 원하는 문자열로 치환합니다."""
    if text:
        # NULL 바이트 및 제어 문자 제거
        text = re.sub(r'[\x00-\x1F\x7F]', '', text)
        # 문자열 치환 추가
        text = text.replace('â', ' ').replace('â', '○').replace('â', '✕')
        return text.strip()
    return ""

def add_table_of_contents(doc):
    """문서의 맨 앞에 목차를 추가합니다."""
    # 새로운 단락 생성
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # fldSimple 요소 생성
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')

    # run 요소를 fldSimple에 추가
    fldSimple.append(run._r)

    # 문서의 가장 앞에 fldSimple 요소를 추가
    doc.element.body.insert(0, fldSimple)

def create_code_style(doc):
    """문서에 코드 블록 스타일을 생성합니다."""
    styles = doc.styles
    if 'Code' not in styles:
        style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Consolas'
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.right_indent = Inches(0.5)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'EDEDED')  # 배경색 설정
        style.paragraph_format.element.get_or_add_pPr().append(shading_elm)

def fetch_and_convert(urls):
    """주어진 URL 목록에서 내용을 가져와 Word 문서로 변환합니다."""
    doc = Document()
    create_code_style(doc)  # 코드 스타일 생성

    for url in urls:
        print(f"Processing {url}...")
        response = requests.get(url)
        response.encoding = 'utf-8'  # 인코딩 설정
        soup = BeautifulSoup(response.text, 'html.parser')

        # theme-doc-markdown 요소 찾기
        content = soup.find(class_='theme-doc-markdown')
        
        # content를 raw HTML 형태로 txt 파일로 저장
        if content:
            with open('content.txt', 'w', encoding='utf-8') as f:
                f.write(str(content))
                f.write('\n\n')

        if content:
            # 콘텐츠를 재귀적으로 순회하여 모든 요소를 처리
            parse_element(content, doc)

    doc.save('output.docx')  # 최종 DOCX 파일 저장

def add_hyperlink(paragraph, text, url):
    """단락에 하이퍼링크를 추가하고, 색상과 밑줄 스타일을 직접 설정합니다."""
    # 하이퍼링크 URL 수정
    if not url.startswith("http"):
        myurl = "https://docs.whatap.io" + url
    else:
        myurl = url

    # 하이퍼링크 ID를 생성합니다.
    part = paragraph.part
    r_id = part.relate_to(myurl, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 하이퍼링크 XML 요소를 생성합니다.
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 하이퍼링크에 포함될 텍스트 Run을 생성합니다.
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 밑줄 설정
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # 색상 설정 (파란색)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # rPr을 run에 추가
    new_run.append(rPr)

    # 텍스트를 추가합니다.
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    # run을 하이퍼링크에 추가
    hyperlink.append(new_run)

    # 단락에 하이퍼링크를 추가합니다.
    paragraph._p.append(hyperlink)

def parse_element(element, doc, parent_style=None):
    """HTML 요소를 재귀적으로 순회하여 Word 문서에 추가합니다."""
    if isinstance(element, NavigableString):
        # 텍스트 노드 처리
        text = clean_text(str(element))
        if text:
            # 현재 단락이 있으면 이어서 추가
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.paragraphs[-1].add_run(text)
            else:
                doc.add_paragraph(text)
        return

    # `<article>` 요소 중 클래스에 `margin-bottom--lg`가 포함된 경우 제외
    if element.name == 'article' and 'margin-bottom--lg' in element.get('class', []):
        return

    if element.name == 'hr':
        # hr 요소는 처리하지 않고 건너뜁니다.
        return

    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5']:
        level = int(element.name[1])
        heading_text = clean_text(element.get_text(strip=True))
        doc.add_heading(heading_text, level=level)
    elif element.name == 'p':
        # 새로운 단락 생성
        paragraph = doc.add_paragraph()
        parse_p_element(element, paragraph)
    elif element.name == 'ul':
        # 리스트 처리
        for li in element.find_all('li', recursive=False):
            parse_element(li, doc, parent_style='ListBullet')
    elif element.name == 'ol':
        # 순서 있는 리스트 처리
        for li in element.find_all('li', recursive=False):
            parse_element(li, doc, parent_style='ListNumber')
    elif element.name == 'li':
        # 리스트 아이템 처리
        paragraph = doc.add_paragraph(style=parent_style or 'ListBullet')
        for child in element.children:
            if isinstance(child, NavigableString):
                text = clean_text(str(child))
                if text:
                    paragraph.add_run(text)
            elif child.name == 'p':
                # li 안의 p 요소 처리
                parse_p_element(child, paragraph)
            elif child.name == 'strong':
                # Bold text for <strong> tags
                text = clean_text(child.get_text())
                if text:
                    # 이전 텍스트와의 공백 처리
                    if paragraph.text and not paragraph.text.endswith(' '):
                        paragraph.add_run(' ')
                    paragraph.add_run(text).bold = True
            elif child.name == 'a':
                text = clean_text(child.get_text())
                href = child.get('href', '')
                if text and href:
                    # 이전 텍스트와의 공백 처리
                    add_hyperlink(paragraph, text, href)
            else:
                # 기타 요소 처리
                parse_element(child, doc)
    elif element.name == 'table':
        # 테이블 처리
        process_table(element, doc)
    elif element.name == 'div' and 'theme-admonition' in element.get('class', []):
        # 경고 메시지 처리
        admonition_title = element.find(class_=lambda c: c and 'admonitionHeading' in c)
        admonition_content = element.find(class_=lambda c: c and 'admonitionContent' in c)
        if admonition_title and admonition_content:
            # 제목 처리
            doc.add_paragraph(f"[{clean_text(admonition_title.get_text(strip=True))}]")
            # 내용 처리
            for child in admonition_content.children:
                parse_element(child, doc)
    elif element.name == 'details':
        # details 태그 처리
        summary = element.find('summary')
        if summary:
            summary_text = clean_text(summary.get_text(strip=True))
            doc.add_paragraph(f"Details: {summary_text}", style='Quote')  # 스타일 변경
            for child in element.children:
                if child != summary:
                    parse_element(child, doc)
    elif element.name == 'img':
        # 이미지 처리
        process_image(element, doc)
    elif element.name == 'div' and 'theme-code-block' in element.get('class', []):
        # 코드 블록 처리
        process_code_block(element, doc)
    else:
        # 기타 요소의 경우 하위 요소를 재귀적으로 처리
        for child in element.children:
            parse_element(child, doc)

def parse_p_element(element, paragraph):
    """<p> 태그의 내용을 처리하여 단락에 추가합니다."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = clean_text(str(child))
            if text:
                # 이전 텍스트와의 공백 처리
                if paragraph.text and not paragraph.text.endswith(' '):
                    paragraph.add_run(' ')
                paragraph.add_run(text)
        elif child.name == 'strong':
            # Bold text for <strong> tags
            text = clean_text(child.get_text())
            if text:
                # 이전 텍스트와의 공백 처리
                if paragraph.text and not paragraph.text.endswith(' '):
                    paragraph.add_run(' ')
                paragraph.add_run(text).bold = True
        elif child.name == 'a':
            text = clean_text(child.get_text())
            href = child.get('href', '')
            if text and href:
                # 이전 텍스트와의 공백 처리
                if paragraph.text and not paragraph.text.endswith(' '):
                    paragraph.add_run(' ')
                # 하이퍼링크 추가
                add_hyperlink(paragraph, text, href)
            else:
                # 하이퍼링크가 없으면 텍스트만 추가
                paragraph.add_run(text)
        elif child.name == 'span' and 'uitext' in child.get('class', []):
            # Apply blue color for <span class="uitext"> elements
            text = clean_text(child.get_text())
            if text:
                # 이전 텍스트와의 공백 처리
                if paragraph.text and not paragraph.text.endswith(' '):
                    paragraph.add_run(' ')
                run = paragraph.add_run(text)
                rPr = run._element.get_or_add_rPr()
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '0000FF')
                rPr.append(color)
        elif child.name == 'img':
            # 이미지 처리
            process_image(child, paragraph)
        else:
            # 다른 태그가 있을 경우 재귀적으로 처리
            parse_p_element(child, paragraph)

def process_code_block(element, doc):
    """코드 블록 요소를 처리하여 문서에 추가합니다."""
    pre_element = element.find('pre')
    if pre_element:
        code_text = pre_element.get_text()
        if code_text:
            # 코드 블록을 문서에 추가
            paragraph = doc.add_paragraph(style='Code')
            paragraph.add_run(code_text)

def process_image(element, container):
    """이미지 요소를 처리하여 문서에 추가합니다."""
    img_src = element.get('src')
    if img_src:
        try:
            if img_src.startswith('data:image'):
                # Base64로 인코딩된 이미지 처리
                header, encoded = img_src.split(',', 1)
                img_data = base64.b64decode(encoded)
            else:
                if not img_src.startswith('http'):
                    img_src = 'https://docs.whatap.io' + img_src
                img_data = requests.get(img_src).content

            img = PILImage.open(BytesIO(img_data))
            img.save("temp_image.png")

            # container의 타입에 따라 이미지 추가 방법 결정
            if hasattr(container, 'add_picture'):
                # Document 또는 Paragraph 객체
                container.add_picture("temp_image.png", width=Inches(5))
            elif hasattr(container, 'add_run'):
                # Paragraph 객체
                container.add_run().add_picture("temp_image.png", width=Inches(5))
            elif hasattr(container, 'paragraphs'):
                # TableCell 객체
                paragraph = container.paragraphs[0]
                paragraph.add_run().add_picture("temp_image.png", width=Inches(1))
            else:
                print(f"Unknown container type: {type(container)}")
            os.remove("temp_image.png")
        except Exception as e:
            print(f"Failed to load image {img_src}: {e}")

def process_table(table_element, doc):
    """HTML 테이블 요소를 Word 문서에 추가합니다."""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # 테이블의 최대 열 수 계산
    max_cols = 0
    for row in rows:
        col_count = 0
        for cell in row.find_all(['th', 'td']):
            colspan = int(cell.get('colspan', 1))
            col_count += colspan
        if col_count > max_cols:
            max_cols = col_count

    table = doc.add_table(rows=0, cols=max_cols)
    table.style = 'Table Grid'

    grid = []
    row_idx = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        while len(grid) <= row_idx:
            grid.append([])
        grid_row = grid[row_idx]
        col_idx = 0
        for cell in cells:
            # 이미 병합된 위치는 건너뜀
            while col_idx < len(grid_row) and grid_row[col_idx] == 'skip':
                col_idx += 1

            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))

            # 셀에 텍스트 및 이미지 추가
            if len(table.rows) <= row_idx:
                table.add_row()
            while len(table.rows[row_idx].cells) < max_cols:
                table.rows[row_idx].add_cell('')

            table_cell = table.cell(row_idx, col_idx)
            # 셀 내용 처리
            parse_table_cell(cell, table_cell)

            # 셀 병합 처리
            if rowspan > 1 or colspan > 1:
                ensure_table_size(table, row_idx + rowspan - 1, col_idx + colspan - 1)
                merge_cell(table, row_idx, col_idx, rowspan, colspan)

            # 병합된 셀 위치 표시
            for i in range(rowspan):
                while len(grid) <= row_idx + i:
                    grid.append([])
                grid_row = grid[row_idx + i]
                while len(grid_row) < col_idx:
                    grid_row.append(None)
                for j in range(colspan):
                    while len(grid_row) <= col_idx + j:
                        grid_row.append(None)
                    if i == 0 and j == 0:
                        grid_row[col_idx + j] = 'data'
                    else:
                        grid_row[col_idx + j] = 'skip'
            col_idx += colspan
        row_idx += 1

def parse_table_cell(cell_element, table_cell):
    """테이블 셀의 내용을 처리하여 Word 셀에 추가합니다."""
    # 기존 셀의 텍스트를 초기화
    table_cell.text = ''
    # 셀에 Paragraph 추가
    paragraph = table_cell.paragraphs[0]
    for child in cell_element.children:
        if isinstance(child, NavigableString):
            text = clean_text(str(child))
            if text:
                paragraph.add_run(text)
        elif child.name == 'span' and 'uitext' in child.get('class', []):
            # Apply blue color for <span class="uitext"> elements
            text = clean_text(child.get_text())
            if text:
                # 이전 텍스트와의 공백 처리
                if paragraph.text and not paragraph.text.endswith(' '):
                    paragraph.add_run(' ')
                run = paragraph.add_run(text)
                rPr = run._element.get_or_add_rPr()
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '0000FF')
                rPr.append(color)
        elif child.name == 'img':
            # 이미지 처리
            process_image(child, table_cell)
        elif child.name == 'br':
            paragraph.add_run('\n')
        else:
            # 기타 요소 처리
            parse_table_cell(child, table_cell)

def ensure_table_size(table, row_idx, col_idx):
    """테이블의 크기를 조정하여 특정 행과 열에 접근 가능하도록 합니다."""
    while len(table.rows) <= row_idx:
        table.add_row()
    for row in table.rows:
        while len(row.cells) <= col_idx:
            row.add_cell('')

def merge_cell(table, row_idx, col_idx, rowspan, colspan):
    """Word 테이블에서 셀 병합을 처리합니다."""
    start_cell = table.cell(row_idx, col_idx)
    end_row = row_idx + rowspan - 1
    end_col = col_idx + colspan - 1

    ensure_table_size(table, end_row, end_col)

    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Fetch URLs from a file and convert to Word document.')
    parser.add_argument('filename', help='File containing URLs to process')
    args = parser.parse_args()

    filename = args.filename

    # 파일 존재 여부 확인
    if not os.path.isfile(filename):
        print(f"Error: The file '{filename}' does not exist.")
        sys.exit(1)

    # URL 목록 읽기
    with open(filename, 'r', encoding='utf-8') as file:
        urls = [line.strip() for line in file if line.strip()]

    fetch_and_convert(urls)
